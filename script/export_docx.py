#!/usr/bin/env python3
"""
导出《夹缝生长》为 DOCX。

依赖安装：
    pip install markdown python-docx

用法：
    python script/export_docx.py
    python script/export_docx.py -o output/my_book.docx
    python script/export_docx.py --chapter 1
"""

import argparse
import io
import os
import re
import sys
from html.parser import HTMLParser

import markdown
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from PIL import Image

# 中文字体配置
CJK_FONT = "PingFang SC"       # macOS 苹方
LATIN_FONT = "PingFang SC"     # 西文也用苹方保持一致
CODE_FONT = "Menlo"            # 代码字体

# 项目根目录
ROOT_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
BOOK_DIR = os.path.join(ROOT_DIR, "book")
INDEX_FILE = os.path.join(ROOT_DIR, "index.md")
DEFAULT_OUTPUT = os.path.join(ROOT_DIR, "output", "夹缝生长.docx")

# 不需要问题页的文件
SKIP_QUESTION_FILES = {"restart.md", "crack.md", "flomo.md", "acknowledgments.md"}

# 不参与编号的特殊文件
SPECIAL_FILES = {"restart.md", "crack.md", "acknowledgments.md"}


def parse_index():
    """从 index.md 解析书籍结构，返回章节列表。"""
    with open(INDEX_FILE, "r", encoding="utf-8") as f:
        content = f.read()

    title_match = re.search(r"^#\s+(.+)", content, re.MULTILINE)
    book_title = title_match.group(1).strip() if title_match else "夹缝生长"

    chapters = []
    current_part = None
    current_section = None

    for line in content.split("\n"):
        line = line.rstrip()

        part_match = re.match(r"^-\s+(第.+部分：.+|前言：|后记：|第.+部分：)", line)
        if part_match:
            part_text = part_match.group(1).rstrip("：").rstrip(":")
            link_match = re.search(r"\[(.+?)\]\((.+?)\)", line)
            if link_match:
                label = link_match.group(1)
                path = link_match.group(2)
                current_part = part_text
                current_section = None
                chapters.append({
                    "part": current_part,
                    "section": None,
                    "title": label,
                    "file": path,
                    "is_part_header": True,
                })
            else:
                current_part = part_text
                current_section = None
            continue

        section_match = re.match(r"^\s*-\s+(.+?)$", line)
        if section_match and "[" not in line:
            current_section = section_match.group(1).strip()
            continue

        link_match = re.search(r"\[(.+?)\]\((.+?)\)", line)
        if link_match:
            chapters.append({
                "part": current_part,
                "section": current_section,
                "title": link_match.group(1),
                "file": link_match.group(2),
                "is_part_header": False,
            })

    return book_title, chapters


def assign_chapter_numbers(chapters):
    """为每个常规章节分配编号。前言、后记、致谢不编号。"""
    num = 0
    for ch in chapters:
        basename = os.path.basename(ch["file"])
        if basename in SPECIAL_FILES or ch.get("is_part_header"):
            ch["chapter_num"] = None
        elif ch["part"] and re.match(r"第.+部分", ch["part"]):
            num += 1
            ch["chapter_num"] = num
        else:
            ch["chapter_num"] = None


def add_numbering_to_content(content, chapter_num):
    """为章节内容的 h1 和 h2 添加编号。不处理 h3 及以下。"""
    if chapter_num is None:
        return content
    h2_counter = 0
    lines = content.split("\n")
    result = []
    for line in lines:
        if re.match(r"^# ", line):
            line = re.sub(r"^# (.+)", f"# {chapter_num}. \\1", line)
        elif re.match(r"^## ", line):
            h2_counter += 1
            line = re.sub(r"^## (.+)", f"## {chapter_num}.{h2_counter} \\1", line)
        result.append(line)
    return "\n".join(result)


def extract_h2_headings(content):
    """从 markdown 内容中提取 h2 标题列表（原始标题，不含编号）。"""
    headings = []
    for line in content.split("\n"):
        match = re.match(r"^## (.+)", line)
        if match:
            headings.append(match.group(1).strip())
    return headings


def read_chapter(file_path):
    """读取章节 markdown 内容。"""
    full_path = os.path.join(ROOT_DIR, file_path)
    if not os.path.exists(full_path):
        print(f"  警告：文件不存在，跳过 {file_path}", file=sys.stderr)
        return None
    with open(full_path, "r", encoding="utf-8") as f:
        return f.read()


def extract_question(content, file_path):
    """从章节内容中提取引导问题。"""
    basename = os.path.basename(file_path)
    if basename in SKIP_QUESTION_FILES:
        return None, content

    match = re.match(r"^(.+?)\n\n---\n", content, re.DOTALL)
    if match:
        question = match.group(1).strip()
        body = content[match.end():]
        return question, body

    return None, content


# ---------------------------------------------------------------------------
# 图片处理
# ---------------------------------------------------------------------------

def _prepare_image(abs_path):
    """将图片转为 Word 兼容的 JPEG 字节流。

    macOS 上的 PNG（Display P3、alpha 通道等）在 Word 中经常渲染异常，
    统一转为 sRGB 的 JPEG 可以解决兼容性问题。
    """
    img = Image.open(abs_path)
    # 转为 RGB（去掉 alpha 通道），白色背景
    if img.mode in ("RGBA", "LA", "P"):
        background = Image.new("RGB", img.size, (255, 255, 255))
        if img.mode == "P":
            img = img.convert("RGBA")
        background.paste(img, mask=img.split()[-1] if "A" in img.mode else None)
        img = background
    elif img.mode != "RGB":
        img = img.convert("RGB")

    buf = io.BytesIO()
    img.save(buf, format="JPEG", quality=90)
    buf.seek(0)
    return buf


# ---------------------------------------------------------------------------
# HTML -> DOCX 转换器
# ---------------------------------------------------------------------------

class HTMLToDocxConverter(HTMLParser):
    """将 Markdown 转出的 HTML 写入 python-docx Document。"""

    def __init__(self, doc, book_dir, file_path):
        super().__init__()
        self.doc = doc
        self.book_dir = book_dir
        self.file_path = file_path

        # 状态跟踪
        self._paragraph = None
        self._tag_stack = []
        self._list_stack = []  # 嵌套列表跟踪: [("ul"|"ol", counter)]
        self._in_blockquote = False
        self._is_epigraph = False
        self._in_code_block = False
        self._skip_content = False  # 跳过 <img> 等空元素的内容

    def _current_tag(self):
        return self._tag_stack[-1] if self._tag_stack else None

    def _ensure_paragraph(self):
        if self._paragraph is None:
            self._paragraph = self.doc.add_paragraph()
        return self._paragraph

    def _finish_paragraph(self):
        self._paragraph = None

    def handle_starttag(self, tag, attrs):
        attrs_dict = dict(attrs)
        self._tag_stack.append(tag)

        if tag in ("h1", "h2", "h3"):
            self._finish_paragraph()
            level = int(tag[1])
            self._paragraph = self.doc.add_heading(level=level)
            return

        if tag == "p":
            self._finish_paragraph()
            if self._in_blockquote:
                p = self.doc.add_paragraph(style="Quote")
                self._paragraph = p
            else:
                self._paragraph = self.doc.add_paragraph()
            return

        if tag == "blockquote":
            self._in_blockquote = True
            css_class = attrs_dict.get("class", "")
            self._is_epigraph = "epigraph" in css_class
            return

        if tag in ("ul", "ol"):
            self._list_stack.append((tag, 0))
            return

        if tag == "li":
            self._finish_paragraph()
            if self._list_stack:
                list_type, counter = self._list_stack[-1]
                if list_type == "ol":
                    counter += 1
                    self._list_stack[-1] = (list_type, counter)
                indent_level = len(self._list_stack) - 1
                if list_type == "ul":
                    p = self.doc.add_paragraph(style="List Bullet")
                else:
                    p = self.doc.add_paragraph(style="List Number")
                # 设置缩进层级
                if indent_level > 0:
                    p.paragraph_format.left_indent = Cm(1.27 * indent_level)
                self._paragraph = p
            else:
                self._paragraph = self.doc.add_paragraph(style="List Bullet")
            return

        if tag == "pre":
            self._in_code_block = True
            self._finish_paragraph()
            return

        if tag == "code":
            # 内联 code 不需要新段落
            return

        if tag == "img":
            src = attrs_dict.get("src", "")
            self._add_image(src)
            return

        if tag == "hr":
            self._finish_paragraph()
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("* * *")
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            run.font.size = Pt(11)
            self._finish_paragraph()
            return

        if tag in ("table", "thead", "tbody", "tr", "th", "td"):
            # 表格处理比较复杂，这里简化处理：收集文本
            return

        if tag == "br":
            p = self._ensure_paragraph()
            p.add_run("\n")
            return

    def handle_endtag(self, tag):
        if self._tag_stack and self._tag_stack[-1] == tag:
            self._tag_stack.pop()

        if tag in ("h1", "h2", "h3"):
            self._finish_paragraph()
            return

        if tag == "p":
            self._finish_paragraph()
            return

        if tag == "blockquote":
            self._in_blockquote = False
            self._is_epigraph = False
            return

        if tag in ("ul", "ol"):
            if self._list_stack:
                self._list_stack.pop()
            return

        if tag == "li":
            self._finish_paragraph()
            return

        if tag == "pre":
            self._in_code_block = False
            self._finish_paragraph()
            return

    def handle_data(self, data):
        if not data.strip() and not self._in_code_block:
            # 允许保留有意义的空格（如 inline 元素之间）
            if data == " " and self._paragraph is not None:
                self._paragraph.add_run(" ")
            return

        current = self._current_tag()

        # 代码块
        if self._in_code_block:
            self._finish_paragraph()
            p = self.doc.add_paragraph()
            run = p.add_run(data)
            set_run_font(run, font_name=CODE_FONT, cjk_font=CJK_FONT)
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            pf = p.paragraph_format
            pf.left_indent = Cm(1)
            self._paragraph = p
            return

        p = self._ensure_paragraph()

        # 判断是否加粗/斜体
        is_bold = "strong" in self._tag_stack or "b" in self._tag_stack
        is_italic = ("em" in self._tag_stack or "i" in self._tag_stack
                     or self._in_blockquote)
        is_code = "code" in self._tag_stack and not self._in_code_block

        run = p.add_run(data)
        set_run_font(run)

        if is_bold:
            run.bold = True
        if is_italic:
            run.italic = True
        if is_code:
            set_run_font(run, font_name=CODE_FONT, cjk_font=CJK_FONT)
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        if self._in_blockquote:
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    def _add_image(self, src):
        """添加图片到文档。"""
        if src.startswith(("http://", "https://")):
            return
        chapter_dir = os.path.dirname(os.path.join(ROOT_DIR, self.file_path))
        abs_path = os.path.abspath(os.path.join(chapter_dir, src))
        if not os.path.exists(abs_path):
            abs_path = os.path.abspath(os.path.join(self.book_dir, src))
        if os.path.exists(abs_path):
            self._finish_paragraph()
            try:
                # 用 Pillow 转为 JPEG 字节流，确保 Word 兼容
                img_stream = _prepare_image(abs_path)
                self.doc.add_picture(img_stream, width=Cm(12))
                last_paragraph = self.doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as e:
                p = self.doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(f"[图片: {os.path.basename(abs_path)}]")
                set_run_font(run)
                print(f"  警告：图片插入失败 {abs_path}: {e}", file=sys.stderr)
            self._finish_paragraph()


def markdown_to_docx(doc, md_content, book_dir, file_path):
    """将 Markdown 内容转为 DOCX 段落，写入 doc。"""
    md_extensions = ["extra", "toc", "sane_lists", "smarty"]
    html_content = markdown.markdown(md_content, extensions=md_extensions)

    # 标记 epigraph
    html_content = re.sub(
        r'(</h1>\s*)((?:<blockquote>\s*.*?</blockquote>\s*)+)',
        lambda m: m.group(1) + re.sub(
            r'<blockquote>',
            '<blockquote class="epigraph">',
            m.group(2),
        ),
        html_content,
        flags=re.DOTALL,
    )

    converter = HTMLToDocxConverter(doc, book_dir, file_path)
    converter.feed(html_content)


# ---------------------------------------------------------------------------
# 样式配置
# ---------------------------------------------------------------------------

def set_run_font(run, font_name=LATIN_FONT, cjk_font=CJK_FONT):
    """为 run 同时设置西文和中文字体。"""
    run.font.name = font_name
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), cjk_font)


def set_style_font(style, font_name=LATIN_FONT, cjk_font=CJK_FONT):
    """为样式同时设置西文和中文字体。"""
    style.font.name = font_name
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), cjk_font)


def setup_styles(doc):
    """配置文档默认样式。"""
    style = doc.styles["Normal"]
    set_style_font(style)
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.8
    style.paragraph_format.space_after = Pt(6)
    style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # 标题样式
    for level in range(1, 4):
        heading_style = doc.styles[f"Heading {level}"]
        set_style_font(heading_style)
        heading_style.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
        if level == 1:
            heading_style.font.size = Pt(20)
        elif level == 2:
            heading_style.font.size = Pt(15)
        else:
            heading_style.font.size = Pt(13)

    # 引用样式
    quote_style = doc.styles["Quote"]
    set_style_font(quote_style)
    quote_style.font.italic = True
    quote_style.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    quote_style.paragraph_format.left_indent = Cm(1.5)

    # 列表样式
    for list_style_name in ("List Bullet", "List Number"):
        if list_style_name in doc.styles:
            set_style_font(doc.styles[list_style_name])

    # 页面设置
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)


def add_page_break(doc):
    """添加分页符。"""
    doc.add_page_break()


def add_cover_page(doc, title):
    """添加封面页。"""
    for _ in range(12):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    set_run_font(run)
    run.font.size = Pt(32)
    run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    run.bold = True

    add_page_break(doc)


def add_centered_page(doc, text, font_size=Pt(26), color=RGBColor(0x22, 0x22, 0x22)):
    """添加居中文字的独立页面（部分标题、子分类标题等）。"""
    for _ in range(12):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run)
    run.font.size = font_size
    run.font.color.rgb = color
    run.bold = True

    add_page_break(doc)


def add_question_page(doc, question):
    """添加问题页（独立一页，居中斜体）。"""
    for _ in range(12):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(question)
    set_run_font(run)
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    run.italic = True

    add_page_break(doc)


# ---------------------------------------------------------------------------
# 目录生成
# ---------------------------------------------------------------------------

def add_toc_page(doc, chapters):
    """添加目录页，包含章节编号和 h2 子标题。"""
    doc.add_heading("目录", level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER

    seen_parts = set()
    seen_sections = set()

    for ch in chapters:
        chapter_num = ch.get("chapter_num")

        if ch["part"] and ch["part"] not in seen_parts:
            seen_parts.add(ch["part"])
            seen_sections.clear()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(ch["part"])
            set_run_font(run)
            run.font.size = Pt(12)
            run.bold = True
            run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)

        if ch["section"] and ch["section"] not in seen_sections:
            seen_sections.add(ch["section"])
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(ch["section"])
            set_run_font(run)
            run.font.size = Pt(10.5)
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            run.bold = True

        # 章节标题（带编号）
        display_title = f"{chapter_num}. {ch['title']}" if chapter_num else ch["title"]
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(2)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(display_title)
        set_run_font(run)
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

        # h2 子标题
        if chapter_num:
            content = read_chapter(ch["file"])
            if content:
                _, body = extract_question(content, ch["file"])
                h2_headings = extract_h2_headings(body)
                for i, heading in enumerate(h2_headings, 1):
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Cm(3)
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(0)
                    run = p.add_run(f"{chapter_num}.{i} {heading}")
                    set_run_font(run)
                    run.font.size = Pt(9.5)
                    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    add_page_break(doc)


# ---------------------------------------------------------------------------
# 主函数
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(description="导出《夹缝生长》为 DOCX")
    parser.add_argument(
        "-o", "--output",
        default=None,
        help=f"输出路径 (默认: {DEFAULT_OUTPUT})",
    )
    parser.add_argument(
        "--chapter",
        type=int,
        default=None,
        help="导出指定章节（按编号，如 --chapter 1）",
    )
    args = parser.parse_args()

    print("解析目录结构...")
    book_title, chapters = parse_index()
    assign_chapter_numbers(chapters)
    print(f"  书名: {book_title}")
    print(f"  章节数: {len(chapters)}")

    # 按章节过滤
    if args.chapter is not None:
        target = [ch for ch in chapters if ch.get("chapter_num") == args.chapter]
        if not target:
            print(f"错误：找不到第 {args.chapter} 章", file=sys.stderr)
            sys.exit(1)
        chapters = target
        ch = target[0]
        if args.output is None:
            output_dir = os.path.join(ROOT_DIR, "output")
            args.output = os.path.join(output_dir, f"{args.chapter:02d}-{ch['title']}.docx")
        print(f"  导出章节: {args.chapter}. {ch['title']}")
    else:
        if args.output is None:
            args.output = DEFAULT_OUTPUT

    output_dir = os.path.dirname(args.output)
    if output_dir:
        os.makedirs(output_dir, exist_ok=True)

    # 创建文档
    doc = Document()
    setup_styles(doc)

    if args.chapter is None:
        # 封面
        add_cover_page(doc, book_title)

        # 目录
        add_toc_page(doc, chapters)

    print("合并章节内容...")

    seen_parts = set()
    seen_sections = set()

    for ch in chapters:
        content = read_chapter(ch["file"])
        if content is None:
            continue

        chapter_num = ch.get("chapter_num")

        # 部分标题页（单章导出时跳过）
        if args.chapter is None and ch["part"] and ch["part"] not in seen_parts:
            seen_parts.add(ch["part"])
            add_centered_page(doc, ch["part"], font_size=Pt(26))
            seen_sections.clear()

        # 子分类标题页（单章导出时跳过）
        if args.chapter is None and ch["section"] and ch["section"] not in seen_sections:
            seen_sections.add(ch["section"])
            add_centered_page(
                doc, ch["section"],
                font_size=Pt(20),
                color=RGBColor(0x44, 0x44, 0x44),
            )

        # 提取引导问题
        question, content = extract_question(content, ch["file"])

        # 添加编号
        content = add_numbering_to_content(content, chapter_num)

        # 问题页（独立一页，在章节最前面）
        if question:
            add_question_page(doc, question)

        # 章节内容
        markdown_to_docx(doc, content, BOOK_DIR, ch["file"])

        # 章节结束后分页
        add_page_break(doc)

    print("生成 DOCX...")
    doc.save(args.output)
    print(f"完成: {args.output}")


if __name__ == "__main__":
    main()
